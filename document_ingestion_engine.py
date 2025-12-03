# =============================================================================
# DOCUMENT INGESTION ENGINE FOR RAG
# =============================================================================
# Converts PDF/DOCX into semantically enriched chunks for vector databases
# =============================================================================

import json
import re
import hashlib
from typing import List, Dict, Any, Tuple
from collections import defaultdict

# Original extraction imports
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

import pdfplumber
import fitz
import pytesseract
import os
from pathlib import Path
from PIL import Image
import logging

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger("rag_ingestion")


# =============================================================================
# 1. DOCUMENT EXTRACTION (Your Original Code)
# =============================================================================

def iter_block_items(parent):
    if not isinstance(parent, _Document):
        raise ValueError("Unsupported parent for block items")
    for child in parent.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract_from_docx(path):
    doc = Document(path)
    elements = []

    hdr = [p.text for s in doc.sections for p in s.header.paragraphs if p.text.strip()]
    ftr = [p.text for s in doc.sections for p in s.footer.paragraphs if p.text.strip()]
    if hdr or ftr:
        elements.append({"type": "header_footer", "header": hdr, "footer": ftr})

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            elements.append({
                "type": "paragraph",
                "text": block.text,
                "style": block.style.name if block.style else "",
                "indent": getattr(block.paragraph_format.left_indent, 'pt', 0) or 0,
                "first_indent": getattr(block.paragraph_format.first_line_indent, 'pt', 0) or 0,
                "page": None  # DOCX doesn't have page info easily
            })
        else:
            elements.append({
                "type": "table",
                "data": [[cell.text for cell in row.cells] for row in block.rows],
                "page": None
            })
    return elements


def page_needs_ocr(page, min_chars: int = 20) -> bool:
    raw = page.extract_text()
    if not raw:
        return True
    if len(raw.strip()) < min_chars:
        return True
    return False


def extract_text_from_pdf(path):
    elements = []
    with pdfplumber.open(path) as pdf:
        logger.info(f"Opened PDF: {path}, pages={len(pdf.pages)}")

        for i, page in enumerate(pdf.pages, 1):
            elements.append({"type": "page_break", "page": i})
            needs_ocr = page_needs_ocr(page, min_chars=20)

            if not needs_ocr:
                logger.info(f"Page {i}: direct text")
                txt = page.extract_text(layout=True) or ""
                elements.append({
                    "type": "page_text",
                    "mode": "text",
                    "page": i,
                    "text": txt
                })
            else:
                logger.info(f"Page {i}: OCR")
                img = page.to_image(resolution=300).original.convert("RGB")
                ocr_txt = pytesseract.image_to_string(img)
                elements.append({
                    "type": "page_text",
                    "mode": "ocr",
                    "page": i,
                    "text": ocr_txt
                })
    return elements


def extract_images_from_pdf(path, images_dir="images"):
    doc = fitz.open(path)
    images_dir = Path(images_dir)
    images_dir.mkdir(parents=True, exist_ok=True)

    image_elements = []
    for page_index in range(len(doc)):
        page = doc.load_page(page_index)
        page_num = page_index + 1

        img_list = page.get_images(full=True)
        for img_idx, img in enumerate(img_list, start=1):
            xref = img[0]
            base_image = doc.extract_image(xref)
            img_bytes = base_image["image"]
            ext = base_image["ext"]

            img_name = f"{Path(path).stem}_p{page_num}_img{img_idx}.{ext}"
            img_path = images_dir / img_name
            with open(img_path, "wb") as f:
                f.write(img_bytes)

            # OCR the image for caption
            try:
                pil_img = Image.open(img_path)
                ocr_text = pytesseract.image_to_string(pil_img).strip()
            except Exception as e:
                logger.warning(f"OCR failed for {img_path}: {e}")
                ocr_text = ""

            bbox = None
            try:
                rects = page.get_image_bbox(xref)
                if isinstance(rects, list) and rects:
                    rect = rects[0]
                else:
                    rect = rects
                bbox = [rect.x0, rect.y0, rect.x1, rect.y1]
            except Exception:
                pass

            image_elements.append({
                "type": "image",
                "page": page_num,
                "bbox": bbox,
                "image_path": str(img_path),
                "ocr_text": ocr_text,
                "index": img_idx
            })

    doc.close()
    return image_elements


def extract_tables_from_pdf(path):
    """Extract tables with pdfplumber"""
    tables = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            page_tables = page.extract_tables()
            for idx, table_data in enumerate(page_tables, 1):
                if table_data:
                    tables.append({
                        "type": "table",
                        "page": i,
                        "data": table_data,
                        "index": idx
                    })
    return tables


def extract_from_pdf(path):
    text_elems = extract_text_from_pdf(path)
    img_elems = extract_images_from_pdf(path)
    table_elems = extract_tables_from_pdf(path)
    return text_elems + img_elems + table_elems


# =============================================================================
# 2. TEXT CLEANING & NORMALIZATION
# =============================================================================

def clean_text(text: str) -> str:
    """Remove artifacts, normalize spacing, preserve structure"""
    if not text:
        return ""
    
    # Remove page numbers (common patterns)
    text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
    text = re.sub(r'\nPage \d+( of \d+)?\n', '\n', text, flags=re.IGNORECASE)
    
    # Remove excessive whitespace
    text = re.sub(r' +', ' ', text)  # Multiple spaces -> single
    text = re.sub(r'\n\n\n+', '\n\n', text)  # Max 2 newlines
    
    # Remove broken words (hyphenation)
    text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
    
    return text.strip()


def is_header_footer(text: str) -> bool:
    """Detect common header/footer patterns"""
    patterns = [
        r'^\d+$',  # Just a number
        r'^Page \d+',
        r'©.*\d{4}',  # Copyright
        r'^confidential$',
        r'^proprietary',
    ]
    text_lower = text.lower().strip()
    if len(text_lower) < 5:
        return True
    for pattern in patterns:
        if re.search(pattern, text_lower):
            return True
    return False


# =============================================================================
# 3. TABLE FORMATTING
# =============================================================================

def table_to_markdown(table_data: List[List[str]]) -> str:
    """Convert table to Markdown format"""
    if not table_data or len(table_data) < 2:
        return ""
    
    # Clean cells
    cleaned = [[str(cell or "").strip() for cell in row] for row in table_data]
    
    # Calculate column widths
    col_widths = [max(len(row[i]) if i < len(row) else 0 for row in cleaned) 
                  for i in range(max(len(row) for row in cleaned))]
    
    lines = []
    
    # Header
    header = cleaned[0]
    lines.append("| " + " | ".join(cell.ljust(w) for cell, w in zip(header, col_widths)) + " |")
    lines.append("|" + "|".join("-" * (w + 2) for w in col_widths) + "|")
    
    # Rows
    for row in cleaned[1:]:
        padded = [row[i].ljust(col_widths[i]) if i < len(row) else " " * col_widths[i] 
                  for i in range(len(col_widths))]
        lines.append("| " + " | ".join(padded) + " |")
    
    return "\n".join(lines)


def summarize_table(table_data: List[List[str]]) -> str:
    """Generate 2-3 line table summary"""
    if not table_data:
        return "Empty table"
    
    rows = len(table_data)
    cols = len(table_data[0]) if table_data else 0
    headers = table_data[0] if table_data else []
    
    summary = f"Table with {rows} rows and {cols} columns. "
    if headers:
        summary += f"Columns: {', '.join(str(h) for h in headers[:5])}{'...' if len(headers) > 5 else ''}."
    
    return summary


# =============================================================================
# 4. IMAGE CAPTION GENERATION
# =============================================================================

def generate_image_caption(img_element: Dict[str, Any]) -> str:
    """Generate detailed image caption"""
    page = img_element.get("page", "?")
    idx = img_element.get("index", "?")
    ocr_text = img_element.get("ocr_text", "").strip()
    bbox = img_element.get("bbox")
    
    caption_parts = []
    
    # Basic description
    caption_parts.append(f"Image on page {page}")
    
    # Position info
    if bbox:
        caption_parts.append(f"located at coordinates {bbox}")
    
    # OCR text
    if ocr_text:
        caption_parts.append(f"containing text: '{ocr_text[:100]}{'...' if len(ocr_text) > 100 else ''}'")
    else:
        caption_parts.append("(no readable text detected)")
    
    caption = ", ".join(caption_parts) + "."
    
    return caption


def format_image_block(img_element: Dict[str, Any]) -> str:
    """Format image as a text block for chunking"""
    page = img_element.get("page", "?")
    idx = img_element.get("index", "?")
    caption = generate_image_caption(img_element)
    ocr_text = img_element.get("ocr_text", "").strip()
    
    lines = [
        f"Image_{page}_{idx}:",
        f"Caption: {caption}"
    ]
    
    if ocr_text:
        lines.append(f"OCR: {ocr_text}")
    
    return "\n".join(lines)


# =============================================================================
# 5. TOKENIZATION (Approximate)
# =============================================================================

def estimate_tokens(text: str) -> int:
    """Rough token estimation: ~4 chars per token"""
    return len(text) // 4


# =============================================================================
# 6. SEMANTIC CHUNKING ENGINE
# =============================================================================

class SemanticChunker:
    def __init__(self, min_tokens=200, max_tokens=350):
        self.min_tokens = min_tokens
        self.max_tokens = max_tokens
        self.chunks = []
    
    def chunk_elements(self, elements: List[Dict], filename: str) -> List[Dict]:
        """Main chunking pipeline"""
        self.chunks = []
        current_chunk = {
            "content": [],
            "pages": set(),
            "content_types": set(),
            "tokens": 0
        }
        
        for elem in elements:
            elem_type = elem.get("type")
            
            # Skip headers/footers
            if elem_type == "header_footer":
                continue
            
            # Page break - track page number
            if elem_type == "page_break":
                continue
            
            # Process different element types
            if elem_type == "paragraph":
                text = clean_text(elem.get("text", ""))
                if not text or is_header_footer(text):
                    continue
                
                self._add_to_chunk(current_chunk, text, elem.get("page"), "text")
                
                # Break on headings
                style = elem.get("style", "")
                if "Heading" in style or style.lower() == "title":
                    current_chunk = self._finalize_chunk(current_chunk, filename)
            
            elif elem_type == "page_text":
                text = clean_text(elem.get("text", ""))
                if not text:
                    continue
                
                # Split by paragraphs
                paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
                for para in paragraphs:
                    if is_header_footer(para):
                        continue
                    self._add_to_chunk(current_chunk, para, elem.get("page"), "text")
            
            elif elem_type == "table":
                md_table = table_to_markdown(elem.get("data", []))
                summary = summarize_table(elem.get("data", []))
                
                if md_table:
                    table_block = f"{md_table}\n\nTable_Summary: {summary}"
                    
                    # Tables get their own chunk
                    current_chunk = self._finalize_chunk(current_chunk, filename)
                    self._add_to_chunk(current_chunk, table_block, elem.get("page"), "table")
                    current_chunk = self._finalize_chunk(current_chunk, filename)
            
            elif elem_type == "image":
                img_block = format_image_block(elem)
                
                # Images get their own chunk
                current_chunk = self._finalize_chunk(current_chunk, filename)
                self._add_to_chunk(current_chunk, img_block, elem.get("page"), "image")
                current_chunk = self._finalize_chunk(current_chunk, filename)
        
        # Finalize last chunk
        if current_chunk["content"]:
            self._finalize_chunk(current_chunk, filename)
        
        return self.chunks
    
    def _add_to_chunk(self, chunk: Dict, text: str, page: int, content_type: str):
        """Add text to current chunk, respecting token limits"""
        tokens = estimate_tokens(text)
        
        # If adding this would exceed max, finalize current chunk first
        if chunk["content"] and chunk["tokens"] + tokens > self.max_tokens:
            # Don't finalize here, let caller handle
            pass
        
        chunk["content"].append(text)
        chunk["tokens"] += tokens
        if page:
            chunk["pages"].add(page)
        chunk["content_types"].add(content_type)
        
        # Auto-finalize if we hit max
        if chunk["tokens"] >= self.max_tokens:
            return True  # Signal to finalize
        
        return False
    
    def _finalize_chunk(self, chunk: Dict, filename: str) -> Dict:
        """Convert accumulated chunk into final format"""
        if not chunk["content"] or chunk["tokens"] < 10:
            # Return empty chunk structure
            return {
                "content": [],
                "pages": set(),
                "content_types": set(),
                "tokens": 0
            }
        
        content_text = "\n\n".join(chunk["content"])
        pages = sorted(chunk["pages"]) if chunk["pages"] else [0]
        
        # Determine content type
        types = chunk["content_types"]
        if len(types) > 1:
            content_type = "mixed"
        else:
            content_type = list(types)[0] if types else "text"
        
        # Generate chunk ID
        chunk_id = self._generate_chunk_id(filename, len(self.chunks))
        
        # Create page range
        if pages and pages[0] > 0:
            if len(pages) == 1:
                page_range = f"p{pages[0]}"
            else:
                page_range = f"p{pages[0]}–p{pages[-1]}"
        else:
            page_range = "unknown"
        
        # Generate summary
        summary = self._generate_summary(content_text)
        
        # Extract tags
        tags = self._extract_tags(content_text)
        
        final_chunk = {
            "chunk_id": chunk_id,
            "source": filename,
            "page_range": page_range,
            "content_type": content_type,
            "content": content_text,
            "summary": summary,
            "tags": tags
        }
        
        self.chunks.append(final_chunk)
        
        # Return fresh chunk structure
        return {
            "content": [],
            "pages": set(),
            "content_types": set(),
            "tokens": 0
        }
    
    def _generate_chunk_id(self, filename: str, index: int) -> str:
        """Generate unique chunk ID"""
        base = f"{Path(filename).stem}_{index}"
        hash_obj = hashlib.md5(base.encode())
        return f"{Path(filename).stem}_chunk_{index}_{hash_obj.hexdigest()[:8]}"
    
    def _generate_summary(self, text: str) -> str:
        """Generate 7-25 word semantic summary"""
        # Simple extractive summary: first sentence or key phrases
        sentences = re.split(r'[.!?]\s+', text)
        first_sentence = sentences[0] if sentences else text
        
        words = first_sentence.split()
        if len(words) > 25:
            summary = " ".join(words[:25]) + "..."
        elif len(words) < 7 and len(sentences) > 1:
            summary = " ".join(words + sentences[1].split()[:15])[:25]
        else:
            summary = first_sentence
        
        return summary.strip()
    
    def _extract_tags(self, text: str) -> List[str]:
        """Extract 3-7 semantic tags"""
        tags = []
        
        # Common domain keywords
        domain_keywords = {
            "contract": ["contract", "agreement", "terms", "conditions"],
            "financial": ["payment", "invoice", "financial", "cost", "price"],
            "technical": ["system", "technical", "specification", "requirements"],
            "legal": ["legal", "compliance", "regulation", "liability"],
            "personal": ["name", "address", "contact", "personal"],
            "date": ["date", "deadline", "schedule", "timeline"],
        }
        
        text_lower = text.lower()
        
        # Extract domain tags
        for domain, keywords in domain_keywords.items():
            if any(kw in text_lower for kw in keywords):
                tags.append(domain)
        
        # Extract capitalized entities (names, places)
        entities = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', text)
        entity_tags = list(set(entities[:3]))  # Top 3 unique entities
        tags.extend(entity_tags)
        
        # Extract numbers/dates as context
        if re.search(r'\d{4}', text):
            tags.append("dated")
        if re.search(r'\$\d+', text):
            tags.append("monetary")
        
        # Limit to 3-7 tags
        tags = list(set(tags))[:7]
        if len(tags) < 3:
            # Add generic tags
            if "table" in text.lower():
                tags.append("tabular_data")
            if "image" in text.lower():
                tags.append("visual_content")
            tags.append("general")
        
        return tags[:7]


# =============================================================================
# 7. MAIN INGESTION PIPELINE
# =============================================================================

def ingest_document(input_path: str, output_json: str = None) -> List[Dict]:
    """
    Main ingestion pipeline: extract -> clean -> chunk -> enrich
    
    Args:
        input_path: Path to PDF or DOCX file
        output_json: Optional output path for JSON chunks
    
    Returns:
        List of enriched chunks
    """
    logger.info(f"Starting ingestion: {input_path}")
    
    # Step 1: Extract elements
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".docx":
        elements = extract_from_docx(input_path)
    elif ext == ".pdf":
        elements = extract_from_pdf(input_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    
    logger.info(f"Extracted {len(elements)} elements")
    
    # Step 2: Chunk elements
    chunker = SemanticChunker(min_tokens=200, max_tokens=350)
    chunks = chunker.chunk_elements(elements, Path(input_path).name)
    
    logger.info(f"Created {len(chunks)} chunks")
    
    # Step 3: Save output
    if output_json:
        with open(output_json, "w", encoding="utf-8") as f:
            json.dump(chunks, f, indent=2, ensure_ascii=False)
        logger.info(f"Saved chunks to {output_json}")
    
    return chunks


# =============================================================================
# 8. UTILITY: Print chunks for review
# =============================================================================

def print_chunks(chunks: List[Dict], max_chunks: int = 5):
    """Print chunks in readable format"""
    for i, chunk in enumerate(chunks[:max_chunks], 1):
        print(f"\n{'='*80}")
        print(f"CHUNK {i}/{len(chunks)}")
        print(f"{'='*80}")
        print(f"ID: {chunk['chunk_id']}")
        print(f"Source: {chunk['source']}")
        print(f"Pages: {chunk['page_range']}")
        print(f"Type: {chunk['content_type']}")
        print(f"Summary: {chunk['summary']}")
        print(f"Tags: {', '.join(chunk['tags'])}")
        print(f"\nContent ({estimate_tokens(chunk['content'])} tokens):")
        print("-" * 80)
        print(chunk['content'][:500] + ("..." if len(chunk['content']) > 500 else ""))
    
    if len(chunks) > max_chunks:
        print(f"\n... and {len(chunks) - max_chunks} more chunks")


# =============================================================================
# 9. EXAMPLE USAGE
# =============================================================================

if __name__ == "__main__":
    # Example: Process a document
    input_file = "qatar_test_doc.pdf"
    output_file = "rag_chunks.json"
    
    if os.path.exists(input_file):
        chunks = ingest_document(input_file, output_file)
        print_chunks(chunks, max_chunks=3)
        
        print(f"\n{'='*80}")
        print(f"SUMMARY: Generated {len(chunks)} chunks from {input_file}")
        print(f"Output saved to: {output_file}")
        print(f"{'='*80}")
    else:
        print(f"Error: {input_file} not found")
        print("\nTo use this engine:")
        print("1. Place your PDF/DOCX file in the same directory")
        print("2. Update 'input_file' variable with your filename")
        print("3. Run: python document_ingestion_engine.py")
